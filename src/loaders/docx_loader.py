from __future__ import annotations

import subprocess
from pathlib import Path

from src.loaders.text_loader import make_chunk
from src.preprocess.chunking import chunk_text, normalize_text, section_hint


def load_docx(path: Path, metadata: dict) -> list[dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required to parse .docx files") from exc

    doc = Document(path)
    chunks: list[dict] = []
    section = ""
    para_index = 0
    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if not text:
            continue
        para_index += 1
        hint = section_hint(text)
        if hint:
            section = text[:80]
        chunks.append(
            make_chunk(
                path,
                metadata,
                text,
                f"{metadata.get('doc_id', path.stem)}::p::{para_index}",
                section=section,
            )
        )

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [normalize_text(cell.text) for cell in row.cells]
            values = [v for v in values if v]
            if not values:
                continue
            text = f"表格{table_index} 第{row_index}行：" + "；".join(values)
            chunks.append(
                make_chunk(
                    path,
                    metadata,
                    text,
                    f"{metadata.get('doc_id', path.stem)}::table::{table_index}::{row_index}",
                    section=f"表格{table_index}",
                    extra={"row": str(row_index)},
                )
            )
    return chunks


def load_doc(path: Path, metadata: dict) -> list[dict]:
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            check=True,
            capture_output=True,
            text=True,
        )
        text = result.stdout
    except Exception as exc:
        raise RuntimeError(f"failed to convert .doc with textutil: {path}") from exc

    chunks = []
    for idx, part in enumerate(chunk_text(text, chunk_size=900), start=1):
        chunks.append(
            make_chunk(
                path,
                metadata,
                part,
                f"{metadata.get('doc_id', path.stem)}::doc::{idx}",
                section=section_hint(part),
            )
        )
    return chunks
